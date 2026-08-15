"""Unit tests for src/parser.py."""
import io
from types import SimpleNamespace

import pypdf
import pytest
from docx import Document

from src.parser import extract_text_from_docx, extract_text_from_pdf, parse_resume


# ---------------------------------------------------------------------------
# Helpers — build minimal in-memory fixtures
# ---------------------------------------------------------------------------


def make_pdf_bytes(pages: list[str]) -> bytes:
    """Create a valid in-memory PDF with one text annotation per page."""
    writer = pypdf.PdfWriter()
    for text in pages:
        page = pypdf.PageObject.create_blank_page(width=200, height=200)
        # pypdf blank pages have no content stream by default; add a simple one
        # that encodes the text so extract_text() can read it back.
        page.merge_page(page)  # no-op but ensures a content stream exists
        writer.add_page(page)
        # Use add_annotation to embed text — simplest portable approach:
        # Actually the most reliable way to get extract_text to return text
        # is to write a raw content stream.
        pass

    # Build a PDF whose pages contain extractable text via a content stream.
    # We use pypdf's PdfWriter + compress_content_streams approach:
    buf = io.BytesIO()
    # Rebuild with reportlab-free approach: write raw PDF bytes manually for
    # a single-page PDF with one text object.
    raw = _minimal_text_pdf(pages)
    return raw


def _minimal_text_pdf(pages: list[str]) -> bytes:
    """
    Build a minimal valid multi-page PDF where each page contains one line
    of text that pypdf.extract_text() can retrieve.

    Uses raw PDF syntax — no external library beyond pypdf required.
    """
    # We'll create one content stream per page with a BT...ET text block.
    content_streams = []
    for text in pages:
        # Escape parentheses and backslashes for PDF string syntax
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 50 700 Td ({escaped}) Tj ET".encode()
        content_streams.append(stream)

    # Object numbering (1-based):
    # 1 = catalog, 2 = pages root, 3..N+2 = page objects, N+3..2N+2 = content streams
    n = len(pages)
    obj_count = 2 + n * 2  # catalog + pages + n pages + n content streams

    objects: dict[int, bytes] = {}

    # Content stream objects first (ids: 3+n .. 2+2n)
    for i, cs in enumerate(content_streams):
        oid = 3 + n + i
        objects[oid] = (
            f"{oid} 0 obj\n"
            f"<< /Length {len(cs)} >>\n"
            f"stream\n"
        ).encode() + cs + b"\nendstream\nendobj\n"

    # Page objects (ids: 3 .. 2+n)
    page_refs = " ".join(f"{3 + i} 0 R" for i in range(n))
    for i in range(n):
        oid = 3 + i
        cs_oid = 3 + n + i
        objects[oid] = (
            f"{oid} 0 obj\n"
            f"<< /Type /Page /Parent 2 0 R "
            f"/MediaBox [0 0 612 792] "
            f"/Contents {cs_oid} 0 R "
            f"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>\n"
            f">>\nendobj\n"
        ).encode()

    # Pages root (id 2)
    objects[2] = (
        f"2 0 obj\n"
        f"<< /Type /Pages /Kids [{page_refs}] /Count {n} >>\n"
        f"endobj\n"
    ).encode()

    # Catalog (id 1)
    objects[1] = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"

    # Assemble
    body = b"%PDF-1.4\n"
    offsets: list[int] = []
    for oid in range(1, 2 + n * 2 + 1):
        offsets.append(len(body))
        body += objects[oid]

    xref_offset = len(body)
    xref = f"xref\n0 {obj_count + 1}\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"

    trailer = (
        f"trailer\n<< /Size {obj_count + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    return body + xref.encode() + trailer.encode()


def make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Create a valid in-memory DOCX with the given paragraphs."""
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_uploaded_file(name: str, data: bytes):
    """Return a duck-typed object matching the Streamlit UploadedFile interface."""
    return SimpleNamespace(name=name, read=lambda: data)


# ---------------------------------------------------------------------------
# extract_text_from_pdf
# ---------------------------------------------------------------------------


def test_extract_text_from_pdf_returns_string():
    """A valid single-page PDF with text returns a non-empty string."""
    pdf_bytes = _minimal_text_pdf(["Hello PDF world"])
    result = extract_text_from_pdf(pdf_bytes)
    assert isinstance(result, str)
    assert len(result) > 0


def test_extract_text_from_pdf_contains_page_text():
    """Extracted text contains the content written to the PDF page."""
    pdf_bytes = _minimal_text_pdf(["Python developer resume"])
    result = extract_text_from_pdf(pdf_bytes)
    assert "Python" in result or "python" in result.lower()


def test_extract_text_from_pdf_multipage_order():
    """Text from earlier pages appears before text from later pages."""
    pdf_bytes = _minimal_text_pdf(["First page content", "Second page content"])
    result = extract_text_from_pdf(pdf_bytes)
    assert result.index("First") < result.index("Second")


def test_extract_text_from_pdf_invalid_bytes_raises_value_error():
    """Random bytes that are not a valid PDF raise ValueError."""
    with pytest.raises(ValueError, match="Could not read PDF file"):
        extract_text_from_pdf(b"this is not a pdf")


# ---------------------------------------------------------------------------
# extract_text_from_docx
# ---------------------------------------------------------------------------


def test_extract_text_from_docx_returns_string():
    """A valid DOCX with one paragraph returns a non-empty string."""
    docx_bytes = make_docx_bytes(["Hello DOCX world"])
    result = extract_text_from_docx(docx_bytes)
    assert isinstance(result, str)
    assert "Hello DOCX world" in result


def test_extract_text_from_docx_multiple_paragraphs():
    """All paragraphs are included and joined by newlines."""
    docx_bytes = make_docx_bytes(["First paragraph", "Second paragraph"])
    result = extract_text_from_docx(docx_bytes)
    assert "First paragraph" in result
    assert "Second paragraph" in result
    assert result.index("First") < result.index("Second")


def test_extract_text_from_docx_invalid_bytes_raises_value_error():
    """Random bytes that are not a valid DOCX raise ValueError."""
    with pytest.raises(ValueError, match="Could not read DOCX file"):
        extract_text_from_docx(b"not a docx file")


def test_extract_text_from_docx_empty_document_raises_runtime_error():
    """A DOCX with no paragraph text raises RuntimeError."""
    doc = Document()  # empty — no paragraphs added
    buf = io.BytesIO()
    doc.save(buf)
    empty_docx = buf.getvalue()
    with pytest.raises(RuntimeError, match="No text could be extracted"):
        extract_text_from_docx(empty_docx)


# ---------------------------------------------------------------------------
# parse_resume dispatch
# ---------------------------------------------------------------------------


def test_parse_resume_dispatches_pdf():
    """A .pdf filename is dispatched to the PDF extractor."""
    pdf_bytes = _minimal_text_pdf(["Resume content"])
    f = make_uploaded_file("resume.pdf", pdf_bytes)
    result = parse_resume(f)
    assert isinstance(result, str)
    assert len(result) > 0


def test_parse_resume_dispatches_docx():
    """A .docx filename is dispatched to the DOCX extractor."""
    docx_bytes = make_docx_bytes(["Resume content"])
    f = make_uploaded_file("resume.docx", docx_bytes)
    result = parse_resume(f)
    assert "Resume content" in result


def test_parse_resume_unsupported_extension_raises_value_error():
    """Extensions other than .pdf and .docx raise ValueError."""
    for ext in [".txt", ".jpg", ".png", ".doc", ".xlsx"]:
        f = make_uploaded_file(f"resume{ext}", b"data")
        with pytest.raises(ValueError, match="Unsupported file format"):
            parse_resume(f)


def test_parse_resume_uppercase_extension_accepted():
    """File names with uppercase extensions (.PDF, .DOCX) are accepted."""
    docx_bytes = make_docx_bytes(["Uppercase extension test"])
    f = make_uploaded_file("RESUME.DOCX", docx_bytes)
    result = parse_resume(f)
    assert "Uppercase extension test" in result
